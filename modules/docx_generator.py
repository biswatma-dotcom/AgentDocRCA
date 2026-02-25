import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx(project, version, blocks, change_logs, requirement_sets, block_templates):
    doc = Document()
    
    title = doc.add_heading('Agent Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Project: {project.name}")
    doc.add_paragraph(f"Client: {project.client_name or 'N/A'}")
    doc.add_paragraph(f"Version: {version.version_number}")
    doc.add_paragraph(f"Date: {version.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()
    
    doc.add_heading('Document Content', level=1)
    
    blocks_dict = version.blocks_content if version.blocks_content else {}
    
    template_dict = {str(bt.id): bt for bt in block_templates}
    
    for block in blocks:
        block_id = str(block.id)
        content = blocks_dict.get(block_id, '')
        
        heading = doc.add_heading(block.name, level=2)
        
        if content:
            doc.add_paragraph(content)
        else:
            doc.add_paragraph("[No content]")
    
    if change_logs:
        doc.add_heading('Change Log for this Version', level=1)
        
        for log in change_logs:
            req_set = next((rs for rs in requirement_sets if str(rs.id) == str(log.requirement_set_id)), None)
            block = template_dict.get(str(log.template_block_id), None)
            
            block_heading = doc.add_heading(f"Block: {block.name if block else 'Unknown'}", level=2)
            
            p = doc.add_paragraph()
            p.add_run("Requirement Set: ").bold = True
            p.add_run(req_set.title if req_set else "Unknown")
            
            if log.requirement_bullet_indexes and req_set and req_set.normalized_bullets:
                p = doc.add_paragraph()
                p.add_run("Bullet(s): ").bold = True
                bullets = []
                for idx in log.requirement_bullet_indexes:
                    if idx < len(req_set.normalized_bullets):
                        bullets.append(req_set.normalized_bullets[idx])
                p.add_run("; ".join(bullets))
            
            p = doc.add_paragraph()
            p.add_run("Reason: ").bold = True
            p.add_run(log.reason_why or "N/A")
            
            p = doc.add_paragraph()
            p.add_run("Impact: ").bold = True
            p.add_run(log.impact or "N/A")
            
            doc.add_paragraph()
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_docx_bytes(project, version, blocks, change_logs, requirement_sets, block_templates):
    buffer = generate_docx(project, version, blocks, change_logs, requirement_sets, block_templates)
    return buffer.getvalue()
